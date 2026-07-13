from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]

FILES = [
    "IT_BRIEF_OMS_V2.md",
    "PRD_OMS_V2.md",
    "TECHNICAL_ARCHITECTURE_OMS_V2.md",
]


def convert_markdown_to_docx(src: Path, dst: Path) -> None:
    lines = src.read_text(encoding="utf-8").splitlines()

    doc = Document()
    list_mode = False

    for raw in lines:
        line = raw.rstrip()
        if not line:
            list_mode = False
            continue

        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
            list_mode = False
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
            list_mode = False
        elif line.startswith("# "):
            doc.add_paragraph(line[2:].strip(), style="Heading 1")
            list_mode = False
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            list_mode = True
        elif len(line) > 3 and line[0].isdigit() and line[1:3] == ". ":
            doc.add_paragraph(line[3:].strip(), style="List Number")
            list_mode = True
        elif line.startswith("---"):
            doc.add_page_break()
            list_mode = False
        else:
            if list_mode:
                doc.add_paragraph(line.strip(), style="Body Text")
            else:
                doc.add_paragraph(line.strip())
            list_mode = False

    doc.save(dst)


if __name__ == "__main__":
    generated = []
    for name in FILES:
        src = ROOT / name
        if not src.exists():
            raise FileNotFoundError(f"Missing source file: {src}")
        dst = src.with_suffix(".docx")
        convert_markdown_to_docx(src, dst)
        generated.append(dst)

    for path in generated:
        print(f"Created {path}")
