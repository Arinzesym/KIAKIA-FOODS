from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUTPUT_FILE = Path("CHURCH_GRANT_BUSINESS_PROPOSAL_KIAKIA_FOODS.docx")
LOGO_CANDIDATES = [
    Path("public/exports/logo.png"),
    Path("public/logo.png"),
]


def money(value: int) -> str:
    return f"N{value:,.0f}"


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def main() -> None:
    ask_amount = 2_000_000
    doc = Document()

    # Keep margins practical for printing and easier editing.
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    logo_path = next((path for path in LOGO_CANDIDATES if path.exists()), None)
    if logo_path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(logo_path), width=Inches(2.2))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("BUSINESS PROPOSAL FOR CHURCH GRANT SUPPORT")
    title_run.bold = True
    title_run.font.size = Pt(20)

    biz = doc.add_paragraph()
    biz.alignment = WD_ALIGN_PARAGRAPH.CENTER
    biz_run = biz.add_run("KiaKia Foods")
    biz_run.bold = True
    biz_run.font.size = Pt(15)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Date: {date.today().strftime('%d %B %Y')}\n")
    meta.add_run("Submitted To: CYON, Holy Family Catholic Church, Life Camp, Abuja\n")
    meta.add_run("Prepared by: KiaKia Foods Management Team\n")
    meta.add_run("Location: Abuja, Nigeria")

    doc.add_paragraph()

    add_heading(doc, "1. Executive Summary")
    add_body(
        doc,
        "KiaKia Foods is a growing grocery sourcing and last-mile delivery business serving households and professionals in Abuja. "
        "Our model helps families save time, avoid market stress, and receive affordable groceries at their doorstep. "
        "We are applying for a church grant to strengthen our operations, serve more families, and create dignified jobs for young people in our community.",
    )

    add_heading(doc, "2. Problem Statement")
    add_body(
        doc,
        "Many households in our service area face three recurring challenges: (1) limited time to shop physically, "
        "(2) unstable access to fresh and reliable market goods, and (3) expensive, uncoordinated delivery options. "
        "These challenges affect working families, elderly people, and single parents most.",
    )

    add_heading(doc, "3. Our Solution")
    add_body(
        doc,
        "KiaKia Foods aggregates customer orders, assigns trained runners for market sourcing, and coordinates grouped estate deliveries. "
        "This method lowers delivery cost per household, improves service reliability, and supports transparent order tracking.",
    )

    add_heading(doc, "4. Mission Alignment with CYON and Parish Community Impact")
    add_body(
        doc,
        "This proposal aligns with CYON and parish values of service, stewardship, and youth formation through practical economic opportunity. "
        "The business creates impact through household food access, fair work opportunities, and skill development pathways for young adults in the church community. "
        "Our service also supports vulnerable members through planned and assisted grocery access.",
    )

    add_heading(doc, "5. Target Beneficiaries")
    add_body(
        doc,
        "Primary beneficiaries include working families, church members, elderly households, and low-mobility customers in Abuja estates. "
        "Secondary beneficiaries include runners, dispatch riders, and support staff employed through expanded operations.",
    )

    add_heading(doc, "6. Operational Plan (12-Month)")
    add_body(
        doc,
        "Phase 1 (Months 1-3): Increase sourcing capacity, standardize operations, and onboard additional runners and dispatch riders.\n"
        "Phase 2 (Months 4-8): Expand to additional estates and improve service reliability through route batching.\n"
        "Phase 3 (Months 9-12): Consolidate quality controls, retention programs, and financial reporting.",
    )

    add_heading(doc, "7. Grant Request and Use of Funds")
    add_body(doc, f"Requested Church Grant: {money(ask_amount)}")

    budget_rows = [
        ("Working Capital for Bulk Sourcing", money(700_000), "Stabilize product availability and pricing"),
        ("Logistics and Delivery Expansion", money(450_000), "Rider support, fuel, route execution, service consistency"),
        ("Operations Technology and Tools", money(250_000), "Order management improvements and reporting tools"),
        ("Staff Recruitment and Training", money(250_000), "Onboard and train runners/operations assistants"),
        ("Customer Growth and Outreach", money(200_000), "Community awareness and service adoption"),
        ("Contingency and Risk Buffer", money(150_000), "Manage volatility and emergency needs"),
        ("Total", money(ask_amount), ""),
    ]

    budget_table = doc.add_table(rows=1, cols=3)
    budget_table.style = "Table Grid"
    budget_hdr = budget_table.rows[0].cells
    budget_hdr[0].text = "Budget Item"
    budget_hdr[1].text = "Amount (NGN)"
    budget_hdr[2].text = "Purpose"

    for row in budget_rows:
        cells = budget_table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
        cells[2].text = row[2]

    add_heading(doc, "8. Financial Outlook (Year 1)")
    finance_table = doc.add_table(rows=1, cols=2)
    finance_table.style = "Table Grid"
    finance_hdr = finance_table.rows[0].cells
    finance_hdr[0].text = "Metric"
    finance_hdr[1].text = "Projected Amount (NGN)"

    finance_rows = [
        ("Projected Annual Revenue", money(42_000_000)),
        ("Projected Direct Operating Costs", money(28_500_000)),
        ("Projected Gross Margin", money(13_500_000)),
        ("Projected Net Surplus (after overhead)", money(6_800_000)),
    ]

    for row in finance_rows:
        cells = finance_table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]

    add_heading(doc, "9. Key Performance Indicators")
    kpi_items = [
        "Number of active households served",
        "On-time delivery rate",
        "Cost-per-delivery and delivery margin",
        "Runner productivity and retention",
        "Number of jobs sustained/created",
        "Number of vulnerable households supported",
    ]
    for item in kpi_items:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "10. Governance, Accountability, and Reporting")
    add_body(
        doc,
        "KiaKia Foods commits to transparent stewardship of grant funds. We will submit periodic reports to the church covering fund utilization, "
        "impact metrics, business performance, and lessons learned. We welcome oversight and advisory support from church leadership.",
    )

    add_heading(doc, "11. Risk Management")
    add_body(
        doc,
        "Main risks include market price volatility, logistics delays, and demand fluctuations. Mitigation actions include bulk purchase planning, "
        "estate-based batching, reserve buffers, and ongoing process controls through our order management system.",
    )

    add_heading(doc, "12. Closing Statement")
    add_body(
        doc,
        "With church grant support, KiaKia Foods will scale responsibly, deepen community service, and strengthen household food access while creating meaningful livelihoods. "
        "We respectfully request your partnership to make this expansion possible.",
    )

    doc.add_paragraph("Prepared by:")
    doc.add_paragraph("KiaKia Foods Management Team")

    doc.save(OUTPUT_FILE)
    print(f"Created {OUTPUT_FILE.resolve()}")


if __name__ == "__main__":
    main()
