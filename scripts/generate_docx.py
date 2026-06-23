from docx import Document
from docx.shared import Inches
from pathlib import Path
from svglib.svglib import svg2rlg
from reportlab.graphics import renderPM

src = Path('BUSINESS_TOOLKIT.md')
dst = Path('BUSINESS_TOOLKIT.docx')
logo_svg = Path('public/logo.svg')
logo_png = Path('public/logo.png')
text = src.read_text(encoding='utf-8').splitlines()

logo_png.parent.mkdir(parents=True, exist_ok=True)
if logo_svg.exists():
    drawing = svg2rlg(str(logo_svg))
    renderPM.drawToFile(drawing, str(logo_png), fmt='PNG', backend='PIL')

doc = Document()
if logo_png.exists():
    doc.add_picture(str(logo_png), width=Inches(1.5))
    doc.add_paragraph()

list_mode = False
for line in text:
    line = line.rstrip()
    if not line:
        list_mode = False
        continue
    if line.startswith('### '):
        doc.add_paragraph(line[4:].strip(), style='Heading 3')
        list_mode = False
    elif line.startswith('## '):
        doc.add_paragraph(line[3:].strip(), style='Heading 2')
        list_mode = False
    elif line.startswith('# '):
        doc.add_paragraph(line[2:].strip(), style='Title')
        list_mode = False
    elif line.startswith('- '):
        doc.add_paragraph(line[2:].strip(), style='List Bullet')
        list_mode = True
    elif line[0].isdigit() and line[1:3] == '. ':
        doc.add_paragraph(line[3:].strip(), style='List Number')
        list_mode = True
    elif line.startswith('---'):
        doc.add_page_break()
        list_mode = False
    else:
        if list_mode:
            doc.add_paragraph(line.strip(), style='Body Text')
        else:
            doc.add_paragraph(line.strip())
        list_mode = False

doc.save(dst)
print(f'Created {dst.resolve()}')
